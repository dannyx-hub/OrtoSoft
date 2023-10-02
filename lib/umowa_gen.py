"""
Kod generujacy umowe na podstawie danych klienta
"""
import docx2txt
from docx import Document
from docx.shared import Pt
import re
import datetime


def find_varialable(doc):
    words_in_braces = set()
    for para in doc.paragraphs:
        text = para.text
        matches = re.findall(r'{(.*?)}', text)
        words_in_braces.update(matches)
    result = []
    for x in words_in_braces:
        print(f"x {x}")
        result.append("{"+x+"}")
    return result


def replace_text(doc, old_text, new_txt):
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_text, new_txt)

# dane z bazy powinny przyjsc w slowniku gdzie {kolumna:wartosc}
def generate_dict(key_list, value_list):
    zmienne = {}
    for x in key_list:
        zmienne[x] = ""
    return zmienne






input_file = "/home/danny/PycharmProjects/ortoSoft/umowa_template.docx"
text = docx2txt.process(input_file)

doc = Document(input_file)
result = find_varialable(doc)
print(result)
zmienne = generate_dict(result)
print(zmienne)

# TODO lepszy sposob na uzupelnienie zmiennych

# zmienne['{adres}'] = adres
# zmienne['{imie_nazwisko}'] = imie_nazwisko
# zmienne['{data_umowy}'] = data_umowy
# zmienne["{telefon}"] = telefon
# zmienne['{dane_szyny}'] = dane_szyny
# zmienne['{numer_urzadzenia}'] = numer_urzadzenia
# zmienne['{rok_urzadzenia}'] = rok_urzadzenia
# zmienne['{wartosc_urzadzenia}'] = wartosc_urzadzenia
# zmienne['{data_wydania}'] = data_wydania

for variable, value in zmienne.items():
    replace_text(doc, variable, value)

output_file = "/home/danny/Pobrane/umowa_2.docx"
doc.save(output_file)
print(f"nowy plik docx zostal wygenerowany jako {output_file}")
